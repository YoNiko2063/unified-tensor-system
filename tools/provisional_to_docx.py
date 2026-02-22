#!/usr/bin/env python3
"""
Convert provisional patent .txt files to formatted .docx

Handles:
  - Title / Inventor / Filing Date block
  - ======= Major Section ======= headings
  - Numbered section headings (1. TITLE, 1.1 Subtitle)
  - Equations (4-space indent) -> Courier New
  - Claims with bold CLAIM N. and indented sub-elements (a)(b)...
  - List items (-- bullets and roman-numeral lists)
  - Figure descriptions -> italic
  - Regular paragraphs
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False,
             color=None, mono=False):
    if mono:
        run.font.name = "Courier New"
        run.font.size = Pt(size)
    else:
        run.font.name = name
        run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def para_spacing(p, before=0, after=4):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def add_horizontal_rule(doc):
    """Add a thin horizontal line using paragraph border."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def leading_spaces(line):
    return len(line) - len(line.lstrip(' '))


def is_section_divider(s):
    return re.match(r'^={40,}', s.strip())


def is_numbered_section(s):
    """e.g. '1. STRUCTURAL FEATURE EXTRACTION' - all caps or mostly caps"""
    m = re.match(r'^(\d+)\.\s+([A-Z][A-Z \-/,]+)$', s.strip())
    return m is not None


def is_numbered_subsection(s):
    """e.g. '1.1 Input Representation'"""
    return bool(re.match(r'^\d+\.\d+(\.\d+)?\s+\S', s.strip()))


def is_claim_line(s):
    return bool(re.match(r'^CLAIM\s+\d+\.', s.strip()))


def is_fig_line(s):
    return bool(re.match(r'^FIG\.\s+\d+', s.strip()))


def is_equation_block(line):
    """4+ space indent, non-empty, not a sub-list."""
    sp = leading_spaces(line)
    stripped = line.strip()
    if sp < 4 or not stripped:
        return False
    # sub-list items start with (i), (ii), (iii), (iv), (v) at 2-4 spaces
    # but we treat them as list items only at exactly 2-3 spaces with roman/letter parens
    # at 4+ spaces, these are equation indents
    return True


def is_list_item(s):
    """Lines starting with -- or (i)(ii) etc. at 2-3 space indent."""
    return bool(re.match(r'^--\s', s)) or bool(re.match(r'^\([ivxIVX]+\)\s', s))


def is_claim_sub_element(s):
    """Lines like '(a) text', '(b) text' inside a claim."""
    return bool(re.match(r'^\([a-z]\)\s', s))


# ---------------------------------------------------------------------------
# Core formatter
# ---------------------------------------------------------------------------

def format_provisional(txt_path: Path, docx_path: Path):
    raw = txt_path.read_text(encoding='utf-8')
    lines = raw.splitlines()

    doc = Document()

    # ---- Page setup --------------------------------------------------------
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    # ---- Default style overrides ------------------------------------------
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)

    for lvl, sz, color_rgb in [
        (1, 15, (0, 0, 102)),   # dark blue
        (2, 13, (0, 0,   0)),
        (3, 12, (0, 0,   0)),
    ]:
        st = doc.styles[f'Heading {lvl}']
        st.font.name  = 'Times New Roman'
        st.font.size  = Pt(sz)
        st.font.bold  = True
        st.font.color.rgb = RGBColor(*color_rgb)
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after  = Pt(4)

    # ---- Find where title block ends (first ====) -------------------------
    title_end = 0
    for idx, ln in enumerate(lines):
        if is_section_divider(ln):
            title_end = idx
            break

    # ---- Render title block -----------------------------------------------
    title_parts, inventor_parts, filing_parts = [], [], []
    mode = None
    for ln in lines[:title_end]:
        s = ln.strip()
        if not s:
            continue
        if s == 'PROVISIONAL PATENT APPLICATION':
            p = doc.add_paragraph()
            run = p.add_run('PROVISIONAL PATENT APPLICATION')
            set_font(run, bold=True, size=13)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para_spacing(p, before=0, after=6)
            continue
        if s.startswith('TITLE:'):
            mode = 'title'
            rest = re.sub(r'^TITLE:\s*', '', s)
            if rest:
                title_parts.append(rest)
            continue
        if s.startswith('INVENTOR:'):
            mode = 'inventor'
            rest = re.sub(r'^INVENTOR:\s*', '', s)
            if rest:
                inventor_parts.append(rest)
            continue
        if s.startswith('FILING DATE:'):
            mode = 'filing'
            filing_parts.append(s)
            continue
        if mode == 'title':
            title_parts.append(s)
        elif mode == 'inventor':
            inventor_parts.append(s)
        elif mode == 'filing':
            filing_parts.append(s)

    if title_parts:
        p = doc.add_paragraph()
        run = p.add_run(' '.join(title_parts))
        set_font(run, bold=True, size=16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(p, before=6, after=8)

    if inventor_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Inventor: ')
        set_font(run, bold=True, size=11)
        run2 = p.add_run(' | '.join(inventor_parts))
        set_font(run2, size=11)

    if filing_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(filing_parts[0])
        set_font(run, size=11)
        para_spacing(p, after=10)

    add_horizontal_rule(doc)

    # ---- Main content state machine ----------------------------------------
    i = title_end
    in_divider = False
    pending_section = []

    def flush_section():
        nonlocal pending_section
        if pending_section:
            title = ' '.join(pending_section).strip()
            p = doc.add_heading(title, level=1)
            para_spacing(p, before=14, after=4)
            add_horizontal_rule(doc)
            pending_section = []

    def consume_wrapped_text(start_idx):
        """Collect continuation lines of a paragraph that have no special prefix."""
        parts = [lines[start_idx].strip()]
        j = start_idx + 1
        while j < len(lines):
            nxt = lines[j]
            ns = nxt.strip()
            # stop on blank
            if not ns:
                break
            sp = leading_spaces(nxt)
            # stop on special starters
            if (is_section_divider(ns) or is_claim_line(ns) or is_numbered_section(ns)
                    or is_numbered_subsection(ns) or is_fig_line(ns)):
                break
            # stop if next line is deeply indented (equation) or list item
            if sp >= 4:
                break
            if is_list_item(ns) or is_claim_sub_element(ns):
                break
            parts.append(ns)
            j += 1
        return ' '.join(parts), j

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        sp = leading_spaces(line)

        # --- Section divider ---
        if is_section_divider(line):
            if in_divider:
                # closing divider
                flush_section()
                in_divider = False
            else:
                in_divider = True
            i += 1
            continue

        if in_divider:
            if stripped:
                pending_section.append(stripped)
            i += 1
            continue

        # --- Blank line ---
        if not stripped:
            i += 1
            continue

        # --- Numbered top-level section heading ---
        if is_numbered_section(stripped):
            p = doc.add_heading(stripped, level=2)
            para_spacing(p, before=12, after=3)
            i += 1
            continue

        # --- Numbered subsection heading ---
        if is_numbered_subsection(stripped):
            p = doc.add_heading(stripped, level=3)
            para_spacing(p, before=8, after=2)
            i += 1
            continue

        # --- Figure description ---
        if is_fig_line(stripped):
            # Collect full figure paragraph
            text, i = consume_wrapped_text(i)
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_font(run, italic=True, size=11)
            para_spacing(p, before=4, after=4)
            p.paragraph_format.left_indent = Inches(0.25)
            continue

        # --- CLAIM block ---
        if is_claim_line(stripped):
            # Collect the entire claim (until next blank+CLAIM or blank+====)
            claim_lines_raw = []
            j = i
            while j < len(lines):
                cline = lines[j]
                cs = cline.strip()
                if not cs:
                    # peek: is next non-blank line a new CLAIM or ====?
                    k = j + 1
                    while k < len(lines) and not lines[k].strip():
                        k += 1
                    if k >= len(lines):
                        claim_lines_raw.append(cline)
                        j += 1
                        break
                    nxt = lines[k].strip()
                    if is_claim_line(nxt) or is_section_divider(nxt):
                        break
                    # blank line within claim body - keep
                    claim_lines_raw.append(cline)
                    j += 1
                    continue
                claim_lines_raw.append(cline)
                j += 1

            # Render the claim
            # First line: "CLAIM N.  text..."
            first_raw = claim_lines_raw[0].strip() if claim_lines_raw else ''
            m = re.match(r'^(CLAIM\s+\d+)\.\s+(.*)', first_raw)

            # Opening claim paragraph
            p = doc.add_paragraph()
            para_spacing(p, before=10, after=2)
            if m:
                run_label = p.add_run(m.group(1) + '.  ')
                set_font(run_label, bold=True, size=12)
                run_text = p.add_run(m.group(2))
                set_font(run_text, size=12)
            else:
                p.add_run(first_raw)

            # Remaining lines
            buf_sub = []  # accumulating sub-element lines

            def flush_sub(sub_lines):
                if not sub_lines:
                    return
                # first sub_line starts with (a), (b) etc.
                first = sub_lines[0]
                rest_joined = ' '.join(sub_lines[1:])
                sp2 = doc.add_paragraph()
                sp2.paragraph_format.left_indent  = Inches(0.5)
                sp2.paragraph_format.first_line_indent = Inches(-0.25)
                para_spacing(sp2, before=1, after=1)
                m2 = re.match(r'^\(([a-z])\)\s+(.*)', first)
                if m2:
                    run_lbl = sp2.add_run(f'({m2.group(1)})  ')
                    set_font(run_lbl, bold=False, size=12)
                    body = m2.group(2)
                    if rest_joined:
                        body += ' ' + rest_joined
                    sp2.add_run(body)
                else:
                    sp2.add_run((first + ' ' + rest_joined).strip())

            for cline in claim_lines_raw[1:]:
                cs = cline.strip()
                sp2 = leading_spaces(cline)
                if not cs:
                    flush_sub(buf_sub)
                    buf_sub = []
                    continue
                if is_claim_sub_element(cs):
                    flush_sub(buf_sub)
                    buf_sub = [cs]
                elif buf_sub:
                    # continuation of current sub-element
                    buf_sub.append(cs)
                else:
                    # continuation of main claim text OR a "wherein" clause
                    # attach to the opening paragraph
                    p.add_run(' ' + cs)

            flush_sub(buf_sub)
            i = j
            continue

        # --- Equation / deeply-indented block (4+ spaces) ---
        if sp >= 4:
            eq_lines = []
            j = i
            while j < len(lines):
                eline = lines[j]
                es = eline.strip()
                esp = leading_spaces(eline)
                if not es:
                    # allow one blank line inside equation block
                    if eq_lines and j + 1 < len(lines) and leading_spaces(lines[j+1]) >= 4:
                        eq_lines.append('')
                        j += 1
                        continue
                    break
                if esp < 4:
                    break
                eq_lines.append(es)
                j += 1

            # Remove trailing blanks
            while eq_lines and not eq_lines[-1]:
                eq_lines.pop()

            for eq in eq_lines:
                if not eq:
                    doc.add_paragraph()
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.6)
                para_spacing(p, before=1, after=1)
                run = p.add_run(eq)
                set_font(run, mono=True, size=10)

            i = j
            continue

        # --- List item at 2-3 space indent ---
        if sp in (2, 3) and is_list_item(stripped):
            text, i = consume_wrapped_text(i)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent       = Inches(0.4)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            para_spacing(p, before=1, after=1)
            if stripped.startswith('--'):
                run = p.add_run('\u2014  ' + text[2:].strip())
            else:
                run = p.add_run(text)
            set_font(run, size=12)
            continue

        # --- Roman/letter list item at 2-6 space indent ---
        if sp in (2, 3, 4) and re.match(r'^\([ivxIVXa-z\d]+\)\s', stripped):
            text, i = consume_wrapped_text(i)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent       = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            para_spacing(p, before=1, after=1)
            p.add_run(text)
            continue

        # --- Regular paragraph ---
        text, i = consume_wrapped_text(i)
        if text.strip():
            p = doc.add_paragraph(text)
            para_spacing(p, before=2, after=5)

    doc.save(str(docx_path))
    print(f'  Saved: {docx_path.name}')


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

PROVISIONALS = [
    'provisional_2_regime_classification.txt',
    'provisional_3_code_synthesis.txt',
    'provisional_4_runtime_validation.txt',
    'provisional_5_platform_integration.txt',
]

IP_DIR      = Path('/home/nyoo/projects/unified-tensor-system/ip')
OUTPUT_DIR  = Path('/home/nyoo/Downloads')

if __name__ == '__main__':
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for fname in PROVISIONALS:
        src  = IP_DIR / fname
        dest = OUTPUT_DIR / fname.replace('.txt', '.docx')
        if not src.exists():
            print(f'  SKIP (not found): {src}')
            continue
        print(f'Converting {fname} ...')
        format_provisional(src, dest)
    print('Done.')
